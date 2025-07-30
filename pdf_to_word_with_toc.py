import logging
import sys
from pathlib import Path
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK

def add_table_of_contents(doc: Document) -> None:
    """
    Inserta un campo de Tabla de Contenidos que Word actualizará automáticamente.
    """
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\h \\z \\u')
    p._p.append(fld)


def pdf_to_word_with_toc(
    pdf_path: str,
    output_path: str,
    dpi: int = 300,
    lang: str = 'spa',
    verbose: bool = True
) -> None:
    """
    Convierte un PDF escaneado a Word con índice interactivo.

    Args:
        pdf_path: Ruta al archivo PDF de entrada.
        output_path: Ruta al archivo Word de salida.
    """
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(format='%(asctime)s - %(levelname)s - %(message)s', level=level)

    pdf_file = Path(pdf_path)
    if not pdf_file.exists():
        logging.error(f"El PDF no existe: {pdf_file}")
        return

    logging.info(f"📄 Convirtiendo {pdf_file.name} a imágenes (DPI={dpi})...")
    pages = convert_from_path(str(pdf_file), dpi=dpi)

    doc = Document()
    doc.core_properties.title = f"Conversión de {pdf_file.name}"
    doc.core_properties.author = "pdf_to_word_with_toc"

    # TOC
    doc.add_heading('Tabla de Contenidos', level=1)
    add_table_of_contents(doc)

    for i, page in enumerate(pages, start=1):
        logging.info(f"🔍 Procesando página {i}/{len(pages)}...")
        text = pytesseract.image_to_string(page, lang=lang)

        if i > 1:
            doc.add_page_break()

        hdr = doc.add_heading(f'Página {i}', level=1)
        hdr.alignment = 1
        doc.add_paragraph(text)

    try:
        doc.save(output_path)
        logging.info(f"✅ Guardado: {output_path}")
    except Exception as e:
        logging.error(f"❌ Error al guardar {output_path}: {e}")

if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description='Convertir uno o varios PDFs escaneados a Word con índice.')
    parser.add_argument('pdfs', nargs='+', help='Rutas a los archivos PDF de entrada')
    parser.add_argument('-o', '--output-dir', default='.', help='Directorio de salida para los DOCX')
    parser.add_argument('-d', '--dpi', type=int, default=300, help='DPI para la conversión de imágenes')
    parser.add_argument('-l', '--lang', default='spa', help='Idioma para OCR (ej.: spa, eng)')
    parser.add_argument('-q', '--quiet', action='store_true', help='Modo silencioso sin logs detallados')
    args = parser.parse_args()

    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    for pdf_path in args.pdfs:
        stem = Path(pdf_path).stem
        output_path = out_dir / f"{stem}.docx"
        pdf_to_word_with_toc(
            pdf_path=pdf_path,
            output_path=str(output_path),
            dpi=args.dpi,
            lang=args.lang,
            verbose=not args.quiet
        )

    logging.info("Proceso completado para todos los PDFs.")
